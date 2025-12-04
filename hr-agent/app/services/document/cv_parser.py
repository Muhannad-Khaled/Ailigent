"""CV/Document Parser - Extract text from PDF and DOCX files."""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


async def parse_cv_file(content: bytes, file_type: str) -> str:
    """
    Parse CV file and extract text content.

    Args:
        content: File content as bytes
        file_type: File extension (pdf or docx)

    Returns:
        Extracted text content
    """
    if file_type.lower() == "pdf":
        return extract_text_from_pdf(content)
    elif file_type.lower() == "docx":
        return extract_text_from_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF file."""
    try:
        from PyPDF2 import PdfReader

        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        full_text = "\n\n".join(text_parts)
        return clean_extracted_text(full_text)

    except ImportError:
        logger.error("PyPDF2 not installed. Cannot parse PDF files.")
        raise
    except Exception as e:
        logger.error(f"Failed to parse PDF: {e}")
        raise


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document

        docx_file = io.BytesIO(content)
        doc = Document(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        return clean_extracted_text(full_text)

    except ImportError:
        logger.error("python-docx not installed. Cannot parse DOCX files.")
        raise
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        raise


def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""

    # Remove excessive whitespace
    lines = text.split("\n")
    cleaned_lines = []

    for line in lines:
        # Strip whitespace
        line = line.strip()

        # Skip empty lines but keep paragraph breaks
        if line:
            cleaned_lines.append(line)
        elif cleaned_lines and cleaned_lines[-1] != "":
            cleaned_lines.append("")

    text = "\n".join(cleaned_lines)

    # Remove multiple consecutive blank lines
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")

    return text.strip()


def detect_file_type(filename: str) -> Optional[str]:
    """Detect file type from filename."""
    if not filename:
        return None

    ext = filename.rsplit(".", 1)[-1].lower()
    if ext in ["pdf", "docx", "doc"]:
        return ext
    return None
