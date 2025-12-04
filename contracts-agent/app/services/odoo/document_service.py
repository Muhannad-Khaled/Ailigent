"""Document Service for Odoo Attachments."""

import base64
import logging
from io import BytesIO
from typing import Any, Dict, List, Optional

from app.core.constants import ATTACHMENT_FIELDS, ODOO_ATTACHMENT_MODEL, SUPPORTED_MIMETYPES
from app.core.exceptions import DocumentNotFoundError, DocumentProcessingError
from app.services.odoo.client import OdooClient, get_odoo_client

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for managing contract documents in Odoo."""

    def __init__(self, odoo_client: Optional[OdooClient] = None):
        self.odoo = odoo_client or get_odoo_client()

    def get_document(self, document_id: int) -> Dict[str, Any]:
        """
        Get a document by ID.

        Args:
            document_id: Odoo attachment ID

        Returns:
            Document details

        Raises:
            DocumentNotFoundError: If document not found
        """
        documents = self.odoo.read(
            ODOO_ATTACHMENT_MODEL,
            [document_id],
            ATTACHMENT_FIELDS,
        )

        if not documents:
            raise DocumentNotFoundError(
                f"Document with ID {document_id} not found",
                details={"document_id": document_id},
            )

        return documents[0]

    def get_documents(
        self,
        document_ids: List[int],
        include_content: bool = False,
    ) -> List[Dict[str, Any]]:
        """
        Get multiple documents by IDs.

        Args:
            document_ids: List of Odoo attachment IDs
            include_content: Whether to include file content

        Returns:
            List of document details
        """
        if not document_ids:
            return []

        fields = ATTACHMENT_FIELDS.copy()
        if not include_content and "datas" in fields:
            fields.remove("datas")

        return self.odoo.read(ODOO_ATTACHMENT_MODEL, document_ids, fields)

    def get_document_content(self, document_id: int) -> bytes:
        """
        Get document binary content.

        Args:
            document_id: Odoo attachment ID

        Returns:
            Document content as bytes

        Raises:
            DocumentNotFoundError: If document not found
            DocumentProcessingError: If content cannot be decoded
        """
        document = self.get_document(document_id)

        if not document.get("datas"):
            raise DocumentProcessingError(
                f"Document {document_id} has no content",
                details={"document_id": document_id},
            )

        try:
            return base64.b64decode(document["datas"])
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to decode document content: {e}",
                details={"document_id": document_id, "error": str(e)},
            )

    def extract_text(self, document_id: int) -> str:
        """
        Extract text from a document (PDF or Word).

        Args:
            document_id: Odoo attachment ID

        Returns:
            Extracted text content

        Raises:
            DocumentProcessingError: If extraction fails
        """
        document = self.get_document(document_id)
        mimetype = document.get("mimetype", "")

        if mimetype not in SUPPORTED_MIMETYPES:
            raise DocumentProcessingError(
                f"Unsupported document type: {mimetype}",
                details={"document_id": document_id, "mimetype": mimetype},
            )

        content = self.get_document_content(document_id)

        try:
            if mimetype == "application/pdf":
                return self._extract_pdf_text(content)
            elif "wordprocessingml" in mimetype or mimetype == "application/msword":
                return self._extract_word_text(content)
            else:
                raise DocumentProcessingError(
                    f"No extractor for mimetype: {mimetype}",
                    details={"document_id": document_id, "mimetype": mimetype},
                )
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to extract text: {e}",
                details={"document_id": document_id, "error": str(e)},
            )

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(BytesIO(content))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise DocumentProcessingError(
                f"PDF extraction failed: {e}",
                details={"error": str(e)},
            )

    def _extract_word_text(self, content: bytes) -> str:
        """Extract text from Word document content."""
        try:
            from docx import Document

            doc = Document(BytesIO(content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)
        except Exception as e:
            raise DocumentProcessingError(
                f"Word extraction failed: {e}",
                details={"error": str(e)},
            )

    def search_documents(
        self,
        name_filter: Optional[str] = None,
        mimetype_filter: Optional[List[str]] = None,
        limit: int = 100,
        offset: int = 0,
    ) -> List[Dict[str, Any]]:
        """
        Search for contract documents.

        Args:
            name_filter: Filter by name (contains)
            mimetype_filter: Filter by mimetypes
            limit: Maximum results
            offset: Results offset

        Returns:
            List of matching documents
        """
        domain = []

        if name_filter:
            domain.append(("name", "ilike", name_filter))

        if mimetype_filter:
            domain.append(("mimetype", "in", mimetype_filter))
        else:
            # Default to supported mimetypes
            domain.append(("mimetype", "in", SUPPORTED_MIMETYPES))

        fields = [f for f in ATTACHMENT_FIELDS if f != "datas"]

        return self.odoo.search_read(
            ODOO_ATTACHMENT_MODEL,
            domain,
            fields=fields,
            limit=limit,
            offset=offset,
            order="create_date desc",
        )

    def upload_document(
        self,
        name: str,
        content: bytes,
        mimetype: str,
        description: Optional[str] = None,
    ) -> int:
        """
        Upload a new document to Odoo.

        Args:
            name: Document name
            content: File content as bytes
            mimetype: MIME type
            description: Optional description

        Returns:
            Created attachment ID
        """
        values = {
            "name": name,
            "datas": base64.b64encode(content).decode("utf-8"),
            "mimetype": mimetype,
            "type": "binary",
        }

        if description:
            values["description"] = description

        return self.odoo.create(ODOO_ATTACHMENT_MODEL, values)


def get_document_service() -> DocumentService:
    """Get document service instance."""
    return DocumentService()
